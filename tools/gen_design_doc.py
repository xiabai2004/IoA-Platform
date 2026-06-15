"""Generate B-EP1 design document from official template."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('IoA 分布式网络运维协同平台')
run.bold = True
run.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('设计文档')
run.bold = True
run.size = Pt(14)

notice = doc.add_paragraph()
notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = notice.add_run('（根据作品匿名要求，本材料不含学校、团队及队员身份信息）')
run.italic = True
run.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('所在赛道与赛项：B-EP1').bold = True
doc.add_paragraph()

# ═══ Section 1 ═══
doc.add_heading('一、目标问题与意义价值', level=1)

doc.add_paragraph(
    '应用领域：网络运维自动化、智能体互联网（Internet of Agents, IoA）、AIOps。'
)

doc.add_paragraph(
    '解决的问题：传统网络运维高度依赖人工进行故障诊断与修复，存在三大痛点：'
    '（1）响应延迟高——从告警到修复平均耗时数十分钟；'
    '（2）跨域协同难——不同域的网络设备配置异构，人工排障需反复切换；'
    '（3）经验难复用——运维知识散落在个人脑中，团队无法沉淀为系统能力。'
    '本作品基于智能体互联网（IoA）架构，构建了一个分布式网络运维协同平台，'
    '实现从自然语言意图输入到故障全闭环修复的自动化流程。'
)

doc.add_paragraph(
    '实现的目标与基本功能：'
    '（1）基于大模型（DeepSeek）的自然语言运维指令解析与意图识别；'
    '（2）多智能体协同编排——9 个专业 Agent 覆盖监控、诊断、修复、验证、报告全流程；'
    '（3）MCP（Model Context Protocol）标准化工具协议，实现 Agent 与模拟器之间的协议解耦；'
    '（4）DAG 有向无环图调度引擎，支持拓扑排序、依赖解析、失败重试与并行执行；'
    '（5）图形化仪表盘——拓扑可视化、实时指标、Agent 状态、消息流一站式展示。'
)

doc.add_paragraph(
    '理论意义与应用价值：本作品是对智能体互联网（IoA）理念的工程实践验证，'
    '探索了 MCP/A2A 标准协议在多智能体系统中的应用范式。在产业应用方面，'
    '该平台可部署于数据中心、云计算平台、工业互联网等场景，显著降低网络运维的'
    '人力成本与故障恢复时间（MTTR）。'
)

# ═══ Section 2 ═══
doc.add_heading('二、设计思路与方案', level=1)

doc.add_heading('2.1 总体设计思路', level=2)
doc.add_paragraph(
    '本系统的核心设计理念是"用智能体替代人工，用协议标准化协同"。'
    '整体采用分层架构，自底向上分为四层：'
    '（1）基础设施层——网络模拟器，模拟 4 域（华东/华北/华南/西南）网络拓扑，'
    '支持 6 种故障注入，提供实时指标生成与 WebSocket 推送；'
    '（2）协议标准层——MCP Server 封装 7 个标准化工具，SSE 传输层对外暴露，'
    '同时支持 A2A 协议实现 Agent 间标准化通信；'
    '（3）中间件层——IoAP 消息路由总线，负责任务分发、DAG 调度编排、Agent '
    '注册发现、UCB1 Bandit 语义路由；'
    '（4）Agent 层——9 个专业 Agent（1 orchestrator + 4 monitor + 1 diagnoser '
    '+ 1 repairer + 1 verifier + 1 reporter）。'
)

doc.add_heading('2.2 技术路线', level=2)
doc.add_paragraph(
    '后端：Python 3.12 + FastAPI + Uvicorn；大模型/智能体：DeepSeek API + '
    'LangChain + LangGraph + MCP SDK 1.27；数据存储：SQLite（aiosqlite）；'
    '消息总线：MemoryBus（开发）+ NATS（生产）双模式；前端：原生 JavaScript + '
    'vis-network 拓扑图 + Chart.js 指标图 + Material Icons 图标库。'
)

doc.add_heading('2.3 详细方案设计', level=2)

doc.add_paragraph(
    '2.3.1 Agent 编排流程：用户通过 GUI 输入自然语言运维指令（如"华东地区网络'
    '延迟异常，全流程诊断修复"），orchestrator-agent 调用 DeepSeek LLM 解析意图，'
    '提取目标域和故障类型，匹配预设 DAG 模板（full_remediation 含 '
    'monitor→diagnose→repair→verify→report 5 个节点），提交给 DAG 调度器。'
    '调度器按拓扑排序依次执行，每个节点通过 SemanticRouter + UCB1 Bandit 选择'
    '最优 Agent 实例，Agent 通过 MCP 协议调用模拟器工具完成操作。'
)

doc.add_paragraph(
    '2.3.2 差异化修复策略：针对 6 种故障类型，设计 primary + fallback 双层策略：'
    '（1）link_congestion → traffic_shape → route_switch；'
    '（2）link_outage → link_failover → route_switch；'
    '（3）cpu_overload → restart_service → traffic_shape；'
    '（4）ddos → acl_deploy → traffic_shape；'
    '（5）misconfig → restart_service → traffic_shape；'
    '（6）device_failure → link_failover → traffic_shape。'
    '修复完成后 verifier Agent 复检指标，未达标则触发 diagnose+repair 回退重试。'
)

doc.add_paragraph(
    '2.3.3 MCP 协议集成：MCP Server 运行在独立端口 9000，采用 SSE 传输模式。'
    'Agent 端 AutoToolClient 优先 MCP 调用，不可用时自动降级 HTTP。系统解决了 '
    'Windows 平台 MCP SSE 的 CRLF 行尾兼容性和 HTTP Keep-Alive 连接复用问题。'
)

# ═══ Section 3 ═══
doc.add_heading('三、方案实现', level=1)

doc.add_heading('3.1 技术栈', level=2)
tech_items = [
    '后端框架：Python 3.12 + FastAPI + Uvicorn',
    '大模型/智能体：DeepSeek API + LangChain + LangGraph + MCP SDK 1.27',
    '数据存储：SQLite（aiosqlite 异步驱动）',
    '消息总线：MemoryBus（开发）+ NATS（生产）双模式，WebSocket 实时推送',
    '前端：原生 JavaScript + vis-network + Chart.js + Material Icons',
    '测试框架：pytest + pytest-asyncio（138 个测试用例）',
]
for item in tech_items:
    doc.add_paragraph(item)

doc.add_heading('3.2 系统架构', level=2)
doc.add_paragraph(
    '系统由四个独立服务组成：'
    '（1）IoA Middleware（端口 8000）——核心中间件，集成 IoAP 消息总线、DAG '
    '调度器、Agent 注册中心、A2A 服务器、WebSocket 端点；'
    '（2）Simulator（端口 8001）——网络模拟器，REST API + WebSocket，支持 4 域 '
    '拓扑和 6 种故障注入；'
    '（3）MCP Server（端口 9000）——独立 Starlette 应用，暴露 7 个标准化工具；'
    '（4）GUI Dashboard（/gui）——SPA 单页应用，五面板布局。'
)

doc.add_heading('3.3 Agent 集群', level=2)
agents = [
    'orchestrator-agent（global）：NL 意图解析 + DAG 模板匹配与提交',
    'monitor-agent × 4（east/north/south/west）：域指标采集 + 阈值异常检测',
    'diagnoser-global（global）：LLM 根因分析 + 修复建议生成',
    'repairer-global（global）：差异化修复，primary + fallback 双策略',
    'verifier-global（global）：闭环验证，复检确认修复效果',
    'reporter-global（global）：运维报告自动生成',
]
for a in agents:
    doc.add_paragraph('• ' + a)

# ═══ Section 4 ═══
doc.add_heading('四、运行结果与应用效果', level=1)

doc.add_heading('4.1 核心功能', level=2)
doc.add_paragraph(
    '（1）一键演示：点击按钮自动完成 故障注入→NL 发送→DAG 全流程执行（5 节点，'
    '约 8-12 秒），全程无人干预。'
    '（2）自然语言交互：输入"华东CPU过载请全流程自动诊断修复"，LLM 解析意图，'
    '自动匹配模板启动修复。'
    '（3）多域并发：支持同时对不同域注入不同故障，3 域 3 故障并发测试全部成功。'
)

doc.add_heading('4.2 测试结果', level=2)
tests = [
    '单元测试 + 集成测试：138 个测试用例全部通过（pytest）',
    'MCP 协议测试：7 个工具全部可用，SSE + JSON-RPC 正常',
    '并发稳定性测试：3 域并发故障注入 + 修复，全部成功',
    '端到端测试：故障注入→NL→DAG→修复→验证→报告 全链路闭环通过',
]
for t in tests:
    doc.add_paragraph('• ' + t)

doc.add_heading('4.3 界面展示', level=2)
doc.add_paragraph(
    'GUI 采用暗色指挥中心风格，五大面板布局：拓扑面板（vis-network 实时渲染'
    '4 域拓扑，故障域颜色标记）；指标面板（Chart.js 延迟趋势图，4 域切换）；'
    'Agent 面板（9 个 Agent 在线状态、域归属、负载）；DAG 面板（任务执行记录，'
    '可展开节点详情）；消息流面板（IoAP 协议消息实时展示）。'
    '系统演示视频和在线演示链接随作品一同提交。'
)

# ═══ Section 5 ═══
doc.add_heading('五、创新与特色', level=1)

innovations = [
    ('创新点一：MCP/A2A 双重协议标准化',
     '完整实现 MCP 协议，将模拟器 7 个能力抽象为标准化工具（SSE 传输），'
     '任何兼容 MCP 的外部智能体均可调用。同时支持 A2A 协议实现 Agent 间'
     '结构化任务分发。解决了 Windows 平台 MCP SSE CRLF/Keep-Alive 兼容性问题。'),
    ('创新点二：差异化修复策略引擎',
     '6 种故障 × primary+fallback 双层策略，修复失败自动切换备选方案。'
     'verifier Agent 闭环验证，未达标触发 diagnose+repair 回退，确保真修复。'),
    ('创新点三：UCB1 Bandit 智能路由',
     '基于多臂老虎机在线学习的 Agent 选择机制，根据历史成功率动态调整路由权重，'
     '实现探索-利用平衡，答辩可展示收敛曲线。'),
    ('创新点四：LangGraph 工作流编排',
     '集成 LangGraph 状态图引擎，将 NL 解析→模板匹配→DAG 生成建模为有状态'
     '多步推理链，支持条件分支和错误恢复。'),
    ('创新点五：双总线 + 跨域弹性架构',
     'MemoryBus（零依赖开发）和 NATS Bus（高可用生产）双模式切换。Agent '
     '支持按域水平扩展，全局 Agent 多域共享，兼顾隔离性和资源效率。'),
]
for title, desc in innovations:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# Save
output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '设计文档_B-EP1_IoA平台.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
